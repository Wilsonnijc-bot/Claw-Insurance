from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\桌面\Claw-Insurance\.tmp-doc-cleanlist\清单-内容整理版.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD_FILL = "FFF6D8"
GOLD = "7A5A00"
BLACK = "111827"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size=None, bold=None, color=None, mono=False):
    ascii_font = "Consolas" if mono else "Calibri"
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.insert(0, node)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        set_cant_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, body_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=body_size,
                        bold=(row_index == 0),
                        color=NAVY if row_index == 0 else BLACK,
                    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_paragraph(doc, text, *, bold_prefix=None, after=6, size=11, color=BLACK):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=size, bold=True, color=color)
        second = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(second, size=size, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=8.8, color=BLACK, mono=True)
    return paragraph


def add_callout(doc, title, body, *, warn=False):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, GOLD_FILL if warn else CALLOUT)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10.5, bold=True, color=GOLD if warn else NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.3, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_page_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("SERVER OPERATIONS CHECKLIST")
    set_run_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("内部运维资料  |  第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(footer, "PAGE")
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("服务器与密钥运维清单")
    set_run_font(run, size=26, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("原清单内容整理版｜真实密码与密钥已替换为占位符")
    set_run_font(run, size=11.5, color=MUTED)

    doc.add_heading("一、服务器基础信息", level=1)
    add_table(
        doc,
        ["项目", "原清单记录"],
        [
            ["云服务", "腾讯云轻量应用服务器"],
            ["公网 IP", "43.129.246.127"],
            ["操作系统", "Ubuntu 24.04 LTS"],
            ["SSH 用户名", "ubuntu"],
            ["SSH 密码", "<已移除，请从密码管理器获取>"],
            ["SSH 登录命令", "ssh ubuntu@43.129.246.127"],
            ["防火墙开放端口", "4000/TCP，来源 0.0.0.0/0"],
        ],
        [2700, 6660],
    )
    add_callout(doc, "密码保存", "SSH 密码使用强密码，并保存在密码管理器中。", warn=True)

    doc.add_heading("二、LiteLLM 代理部署路径与配置文件", level=1)
    add_table(
        doc,
        ["项目", "路径或内容"],
        [
            ["项目目录", "/home/ubuntu/litellm-proxy"],
            ["docker-compose.yml", "定义 litellm 和 db 两个服务；映射 4000 端口；挂载 ./config.yaml"],
            ["config.yaml", "配置 model_list、Kimi 模型、API Key 与 Moonshot API Base"],
            [".env", "保存 MOONSHOT_API_KEY、LITELLM_MASTER_KEY、DB_PASSWORD"],
            ["postgres_data", "数据库持久化目录（容器卷）"],
        ],
        [2700, 6660],
    )
    doc.add_heading("config.yaml 原清单结构", level=2)
    add_code(
        doc,
        "model_list:\n"
        "  - model_name: kimi-k2.5\n"
        "    litellm_params:\n"
        "      model: moonshot/kimi-k2.5\n"
        "      api_key: <MOONSHOT_API_KEY>\n"
        "      api_base: https://api.moonshot.cn/v1",
    )

    doc.add_heading("三、关键密钥与密码清单", level=1)
    add_table(
        doc,
        ["名称", "作用", "整理后记录"],
        [
            ["Kimi API Key", "调用 Kimi 官方接口（上游）", "<已移除>"],
            ["LiteLLM 主密钥", "管理代理、生成虚拟密钥、访问管理接口", "<已移除>"],
            ["PostgreSQL 数据库密码", "连接数据库", "<请从密码管理器获取>"],
            ["服务器 SSH 密码", "登录服务器", "<请从密码管理器获取>"],
        ],
        [2200, 4100, 3060],
    )
    add_callout(doc, "敏感信息", "API Key、主密钥、数据库密码和 SSH 密码不在此文档中保存。", warn=True)

    doc.add_heading("四、服务器日常管理命令", level=1)
    command_groups = [
        ("登录服务器", "ssh ubuntu@43.129.246.127"),
        ("进入 LiteLLM 目录", "cd ~/litellm-proxy"),
        ("查看服务状态", "sudo docker compose ps\nsudo docker compose logs litellm --tail 50"),
        ("重启服务（修改配置后）", "sudo docker compose down\nsudo docker compose up -d"),
        ("查看容器资源占用", "sudo docker stats --no-stream\nfree -h"),
        ("进入 LiteLLM 容器内部", "sudo docker exec -it litellm-proxy-litellm-1 bash"),
        ("查看数据库日志", "sudo docker compose logs db --tail 30"),
    ]
    for heading, command in command_groups:
        doc.add_heading(heading, level=2)
        add_code(doc, command)

    doc.add_heading("五、为用户生成虚拟密钥", level=1)
    key_intro = add_paragraph(doc, "每次为一个新用户生成独立虚拟密钥。")
    key_intro.paragraph_format.keep_with_next = True
    doc.add_heading("命令模板", level=2)
    add_code(
        doc,
        "curl -X POST \"http://localhost:4000/key/generate\" \\\n"
        "  -H \"Authorization: Bearer $LITELLM_MASTER_KEY\" \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"alias\": \"用户名\", \"models\": [\"kimi-k2.5\"]}'",
    )
    doc.add_heading("示例：用户名为 zhangsan", level=2)
    add_code(
        doc,
        "curl -X POST \"http://localhost:4000/key/generate\" \\\n"
        "  -H \"Authorization: Bearer $LITELLM_MASTER_KEY\" \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"alias\": \"zhangsan\", \"models\": [\"kimi-k2.5\"]}'",
    )
    add_paragraph(doc, "返回结果：提取 key 字段的值，作为分发给该用户的虚拟密钥。")

    doc.add_heading("六、管理虚拟密钥", level=1)
    doc.add_heading("查看所有虚拟密钥", level=2)
    add_code(
        doc,
        "curl -X GET \"http://localhost:4000/key/list\" \\\n"
        "  -H \"Authorization: Bearer $LITELLM_MASTER_KEY\"",
    )
    doc.add_heading("查看某个虚拟密钥", level=2)
    add_code(
        doc,
        "curl -X GET \"http://localhost:4000/key/info?key=sk-xxx\" \\\n"
        "  -H \"Authorization: Bearer $LITELLM_MASTER_KEY\"",
    )
    doc.add_heading("删除某个虚拟密钥", level=2)
    add_code(
        doc,
        "curl -X POST \"http://localhost:4000/key/delete\" \\\n"
        "  -H \"Authorization: Bearer $LITELLM_MASTER_KEY\" \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"keys\": [\"sk-xxx\"]}'",
    )

    doc.add_heading("七、用户分发信息与配置", level=1)
    add_table(
        doc,
        ["项目", "原清单内容"],
        [
            ["LiteLLM 网关地址", "http://43.129.246.127:4000"],
            ["用户虚拟密钥", "<用户专属虚拟密钥>"],
            ["可用模型名称", "kimi-k2.5"],
        ],
        [2700, 6660],
    )
    doc.add_heading("用户配置步骤", level=2)
    add_table(
        doc,
        ["顺序", "操作"],
        [
            ["1", "克隆仓库"],
            ["2", "复制 config.example.json 为 config.json"],
            ["3", "将 apiKey 字段改为用户虚拟密钥"],
            ["4", "确保 baseUrl 指向 LiteLLM 网关地址"],
            ["5", "运行 docker-compose up -d"],
        ],
        [1100, 8260],
    )
    doc.add_heading("原清单中的其他配置记录", level=2)
    add_table(
        doc,
        ["文件或目录", "记录内容"],
        [
            ["supabaseconfig.json", "将 apiKey 字段替换为 <原值已移除>"],
            ["config.json", "将 apiKey 字段替换为 <用户专属虚拟密钥>"],
            ["secrets/", "放入 double-scholar-487115-b1-075776a1689b.json"],
        ],
        [3000, 6360],
    )

    doc.add_heading("八、生成 DB_PROXY_API_KEY", level=1)
    doc.add_heading("OpenSSL", level=2)
    add_code(doc, "openssl rand -hex 32")
    add_paragraph(doc, "输出为 32 字节随机数据，即 64 个十六进制字符，可作为 DB_PROXY_API_KEY。")

    doc.add_heading("Windows PowerShell", level=2)
    add_code(
        doc,
        "$bytes = [byte[]]::new(32)\n"
        "[System.Security.Cryptography.RandomNumberGenerator]::Create().GetBytes($bytes)\n"
        "$key = [Convert]::ToHexString($bytes).ToLower()\n"
        "Write-Output $key",
    )
    doc.add_heading("Python", level=2)
    add_code(doc, "import secrets\nprint(secrets.token_hex(32))")
    doc.add_heading("Linux / macOS", level=2)
    add_code(doc, "cat /dev/urandom | tr -dc 'a-f0-9' | fold -w 64 | head -n 1")

    doc.add_heading("九、Docker 多架构发布记录", level=1)
    doc.add_heading("执行命令", level=2)
    add_code(
        doc,
        "powershell.exe -NoProfile -ExecutionPolicy Bypass `\n"
        "  -File \".\\scripts\\publish-multiarch.ps1\"",
    )
    add_table(
        doc,
        ["项目", "日志记录"],
        [
            ["Buildx Builder", "claw-multiarch"],
            ["Driver", "docker-container"],
            ["目标平台", "linux/amd64、linux/arm64"],
            ["构建阶段", "45/45 步骤完成"],
            ["推送阶段", "上传 Docker Hub 时出现 broken pipe / closed network connection"],
            ["最终结果", "获取 Docker Hub OAuth token 时 EOF；后端多平台发布失败"],
        ],
        [2700, 6660],
    )

    doc.core_properties.title = "服务器与密钥运维清单"
    doc.core_properties.subject = "原清单内容整理版"
    doc.core_properties.author = ""
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
