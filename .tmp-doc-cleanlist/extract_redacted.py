from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


def redact(text: str) -> str:
    text = re.sub(r"sk-[A-Za-z0-9_-]{12,}", "<REDACTED_API_KEY>", text)
    text = re.sub(
        r"eyJ[A-Za-z0-9_-]{10,}(?:\.[A-Za-z0-9_-]{10,}){1,2}",
        "<REDACTED_JWT>",
        text,
    )
    text = re.sub(
        r"(?i)(password|密码|key|token)(\s*[:：=]\s*)(\S{8,})",
        lambda m: f"{m.group(1)}{m.group(2)}<REDACTED>",
        text,
    )
    return text


path = Path(sys.argv[1])
doc = Document(path)
for index, paragraph in enumerate(doc.paragraphs, start=1):
    value = paragraph.text.strip()
    if value:
        print(f"P{index:03d} [{paragraph.style.name}] {redact(value)}")

for table_index, table in enumerate(doc.tables, start=1):
    print(f"TABLE {table_index}")
    for row_index, row in enumerate(table.rows, start=1):
        values = [redact(cell.text.strip().replace("\n", " / ")) for cell in row.cells]
        print(f"  R{row_index:02d}: " + " || ".join(values))
